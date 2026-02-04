"""Result file matcher to load and index experiment results."""

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Any
import csv
import sqlite3
import pickle

# All format support libraries (built-in dependencies)
import openpyxl
from docx import Document
import yaml
import tabula


@dataclass
class ResultValue:
    """A value from experiment results."""
    
    value: float
    source_file: str
    path: str  # JSON path or CSV column
    model: Optional[str] = None
    dataset: Optional[str] = None
    metric: Optional[str] = None


class ResultMatcher:
    """Load and index experiment result files."""
    
    SUPPORTED_FORMATS = {
        'json': True,
        'csv': True,
        'xlsx': True,
        'xls': True,
        'docx': True,
        'yaml': True,
        'yml': True,
        'sqlite': True,
        'db': True,
        'pdf': True,
        'pkl': True,
        'pickle': True,
    }
    
    def __init__(self):
        self.values: List[ResultValue] = []
        self.index: Dict[float, List[ResultValue]] = {}
    
    def load_directory(self, dirpath: Path) -> int:
        """Load all result files from a directory."""
        dirpath = Path(dirpath)
        count = 0
        
        # Load JSON files
        for json_file in dirpath.rglob('*.json'):
            try:
                self._load_json_file(json_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {json_file}: {e}")
        
        # Load CSV files
        for csv_file in dirpath.rglob('*.csv'):
            try:
                self._load_csv_file(csv_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {csv_file}: {e}")
        
        # Load Excel files
        for xlsx_file in dirpath.rglob('*.xlsx'):
            try:
                self._load_excel_file(xlsx_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {xlsx_file}: {e}")
        for xls_file in dirpath.rglob('*.xls'):
            try:
                self._load_excel_file(xls_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {xls_file}: {e}")
        
        # Load Word files
        for docx_file in dirpath.rglob('*.docx'):
            try:
                self._load_word_file(docx_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {docx_file}: {e}")
        
        # Load YAML files
        for yaml_file in dirpath.rglob('*.yaml'):
            try:
                self._load_yaml_file(yaml_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {yaml_file}: {e}")
        for yml_file in dirpath.rglob('*.yml'):
            try:
                self._load_yaml_file(yml_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {yml_file}: {e}")
        
        # Load SQLite files
        for db_file in dirpath.rglob('*.sqlite'):
            try:
                self._load_sqlite_file(db_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {db_file}: {e}")
        for db_file in dirpath.rglob('*.db'):
            try:
                self._load_sqlite_file(db_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {db_file}: {e}")
        
        # Load PDF files (tables only)
        for pdf_file in dirpath.rglob('*.pdf'):
            try:
                self._load_pdf_file(pdf_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {pdf_file}: {e}")
        
        # Load Pickle files
        for pkl_file in dirpath.rglob('*.pkl'):
            try:
                self._load_pickle_file(pkl_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {pkl_file}: {e}")
        for pickle_file in dirpath.rglob('*.pickle'):
            try:
                self._load_pickle_file(pickle_file)
                count += 1
            except Exception as e:
                print(f"Warning: Could not load {pickle_file}: {e}")
        
        # Build index
        self._build_index()
        
        return count
    
    def _load_json_file(self, filepath: Path):
        """Load values from a JSON file."""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        filename = filepath.stem
        model, dataset = self._parse_filename(filename)
        self._extract_from_dict(data, str(filepath), '', model, dataset)
    
    def _load_csv_file(self, filepath: Path):
        """Load values from a CSV file."""
        with open(filepath, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row_num, row in enumerate(reader):
                for col, val in row.items():
                    try:
                        numeric_val = float(val)
                        self.values.append(ResultValue(
                            value=numeric_val,
                            source_file=str(filepath),
                            path=f"row_{row_num}.{col}",
                            metric=self._guess_metric(col),
                        ))
                    except (ValueError, TypeError):
                        continue
    
    def _load_excel_file(self, filepath: Path):
        """Load values from an Excel file."""
        wb = openpyxl.load_workbook(filepath, data_only=True)
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                for col_idx, cell in enumerate(row):
                    if isinstance(cell, (int, float)) and cell is not None:
                        self.values.append(ResultValue(
                            value=float(cell),
                            source_file=str(filepath),
                            path=f"{sheet_name}.row{row_idx}.col{col_idx}",
                            metric=None,
                        ))
    
    def _load_word_file(self, filepath: Path):
        """Load numeric values from a Word document (tables and text)."""
        import re
        doc = Document(filepath)
        
        # Extract from tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    try:
                        numeric_val = float(text.replace(',', ''))
                        self.values.append(ResultValue(
                            value=numeric_val,
                            source_file=str(filepath),
                            path=f"table{table_idx}.row{row_idx}.col{col_idx}",
                            metric=None,
                        ))
                    except ValueError:
                        pass
        
        # Extract from paragraphs
        number_pattern = re.compile(r'[\d,]+\.?\d*')
        for para_idx, para in enumerate(doc.paragraphs):
            for match in number_pattern.finditer(para.text):
                try:
                    numeric_val = float(match.group().replace(',', ''))
                    if numeric_val != 0:
                        self.values.append(ResultValue(
                            value=numeric_val,
                            source_file=str(filepath),
                            path=f"para{para_idx}",
                            metric=None,
                        ))
                except ValueError:
                    pass
    
    def _load_yaml_file(self, filepath: Path):
        """Load values from a YAML file."""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        if data:
            filename = filepath.stem
            model, dataset = self._parse_filename(filename)
            self._extract_from_dict(data, str(filepath), '', model, dataset)
    
    def _load_sqlite_file(self, filepath: Path):
        """Load numeric values from SQLite database."""
        conn = sqlite3.connect(filepath)
        cursor = conn.cursor()
        
        # Get all tables
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = cursor.fetchall()
        
        for (table_name,) in tables:
            try:
                cursor.execute(f"SELECT * FROM {table_name}")
                columns = [desc[0] for desc in cursor.description]
                
                for row_idx, row in enumerate(cursor.fetchall()):
                    for col_idx, (col_name, value) in enumerate(zip(columns, row)):
                        if isinstance(value, (int, float)) and value is not None:
                            self.values.append(ResultValue(
                                value=float(value),
                                source_file=str(filepath),
                                path=f"{table_name}.{col_name}.row{row_idx}",
                                metric=self._guess_metric(col_name),
                            ))
            except Exception:
                continue
        
        conn.close()
    
    def _load_pdf_file(self, filepath: Path):
        """Load numeric values from PDF tables using tabula-py."""
        # Extract all tables from PDF
        tables = tabula.read_pdf(str(filepath), pages='all', silent=True)
        
        for table_idx, df in enumerate(tables):
            for row_idx in range(len(df)):
                for col_idx, col_name in enumerate(df.columns):
                    val = df.iloc[row_idx, col_idx]
                    try:
                        # Handle string numbers with commas
                        if isinstance(val, str):
                            val = val.replace(',', '')
                        numeric_val = float(val)
                        if numeric_val != 0:
                            self.values.append(ResultValue(
                                value=numeric_val,
                                source_file=str(filepath),
                                path=f"table{table_idx}.{col_name}.row{row_idx}",
                                metric=self._guess_metric(str(col_name)),
                            ))
                    except (ValueError, TypeError):
                        continue
    
    def _load_pickle_file(self, filepath: Path):
        """Load numeric values from a Pickle file (DataFrame or dict)."""
        with open(filepath, 'rb') as f:
            data = pickle.load(f)
        
        # Handle pandas DataFrame
        if hasattr(data, 'iterrows'):  # Duck typing for DataFrame
            for row_idx, row in data.iterrows():
                for col_name in data.columns:
                    val = row[col_name]
                    try:
                        numeric_val = float(val)
                        self.values.append(ResultValue(
                            value=numeric_val,
                            source_file=str(filepath),
                            path=f"{col_name}.row{row_idx}",
                            metric=self._guess_metric(str(col_name)),
                        ))
                    except (ValueError, TypeError):
                        continue
        # Handle dict/list
        elif isinstance(data, (dict, list)):
            filename = filepath.stem
            model, dataset = self._parse_filename(filename)
            self._extract_from_dict(data, str(filepath), '', model, dataset)
        # Handle single numeric value
        elif isinstance(data, (int, float)):
            self.values.append(ResultValue(
                value=float(data),
                source_file=str(filepath),
                path='value',
                metric=None,
            ))
    
    def _extract_from_dict(
        self, 
        data: Any, 
        source_file: str, 
        path: str,
        model: Optional[str],
        dataset: Optional[str]
    ):
        """Recursively extract numeric values from a dictionary."""
        if isinstance(data, dict):
            for key, value in data.items():
                new_path = f"{path}.{key}" if path else key
                self._extract_from_dict(value, source_file, new_path, model, dataset)
        elif isinstance(data, list):
            if len(data) > 100:
                return
            for i, item in enumerate(data):
                self._extract_from_dict(item, source_file, f"{path}[{i}]", model, dataset)
        elif isinstance(data, (int, float)):
            metric = self._guess_metric(path)
            self.values.append(ResultValue(
                value=float(data),
                source_file=source_file,
                path=path,
                model=model,
                dataset=dataset,
                metric=metric,
            ))
    
    def _parse_filename(self, filename: str) -> tuple:
        """Extract model and dataset from filename like 'xgboost_etth1'."""
        parts = filename.lower().replace('-', '_').split('_')
        
        model = None
        dataset = None
        
        models = ['xgboost', 'arima', 'chronos', 'moirai', 'dlinear', 'patchtst', 'timesfm']
        datasets = ['etth1', 'etth2', 'exchange', 'traffic', 'weather']
        
        for part in parts:
            if any(m in part for m in models):
                model = part
            if any(d in part for d in datasets):
                dataset = part
        
        return model, dataset
    
    def _guess_metric(self, path: str) -> Optional[str]:
        """Guess the metric type from the JSON path."""
        path_lower = path.lower()
        
        if 'mae' in path_lower:
            return 'mae'
        elif 'rmse' in path_lower:
            return 'rmse'
        elif 'smape' in path_lower:
            return 'smape'
        elif 'latency' in path_lower:
            return 'latency'
        elif 'vram' in path_lower:
            return 'vram'
        
        return None
    
    def _build_index(self):
        """Build an index for fast lookup."""
        self.index.clear()
        for rv in self.values:
            key = round(rv.value, 6)
            if key not in self.index:
                self.index[key] = []
            self.index[key].append(rv)
    
    def find_matches(
        self, 
        target: float, 
        tolerance_pct: float = 1.0
    ) -> List[ResultValue]:
        """Find result values that match the target within tolerance."""
        matches = []
        
        for rv in self.values:
            if rv.value == 0:
                continue
            
            diff_pct = abs(rv.value - target) / abs(rv.value) * 100
            
            if diff_pct <= tolerance_pct:
                matches.append(rv)
        
        return matches
    
    def find_exact(self, target: float) -> List[ResultValue]:
        """Find exact matches (within floating point tolerance)."""
        key = round(target, 6)
        return self.index.get(key, [])


def load_results(dirpath: Path) -> ResultMatcher:
    """Convenience function to load results from a directory."""
    matcher = ResultMatcher()
    matcher.load_directory(dirpath)
    return matcher

